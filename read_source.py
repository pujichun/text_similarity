import abc

import docx
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LAParams, LTTextBoxHorizontal
from pdfminer.pdfinterp import (PDFPageInterpreter, PDFResourceManager,
                                PDFTextExtractionNotAllowed)
from pdfminer.pdfparser import PDFDocument, PDFParser


# define a reader interface
class Reader(metaclass=abc.ABCMeta):
    def __init__(self, path: str) -> None:
        self.path = path

    @abc.abstractmethod
    def read(self) -> str:
        pass


class TxtRadeer(Reader):
    def read(self) -> str:
        f = open(file=self.path, mode="r", encoding="utf8")
        t = f.read()
        return t.replace(" ", "")


class WordReader(Reader):
    def read(self) -> str:
        document = docx.Document(self.path)
        w = ""
        for page in document.paragraphs:
            w += page.text
        return w.replace(" ", "")


class PDFReader(Reader):
    def read(self) -> str:
        source = ""
        with open(file=self.path, mode="rb") as f:
            # create a PDFParser object
            parser = PDFParser(f)
            # create PDFDocument object
            pdf_file = PDFDocument()
            # link analyzer and document object
            parser.set_document(pdf_file)
            pdf_file.set_parser(parser)
            # provide initialization password
            pdf_file.initialize()
            # Check whether the document provides txt conversion
        if not pdf_file.is_extractable:
            raise PDFTextExtractionNotAllowed
        else:
            # Interpret data
            # manage data
            manager = PDFResourceManager()
            # Create a PDF device object of LAParams
            layout_params = LAParams()
            device = PDFPageAggregator(manager, laparams=layout_params)
            # PDFInterpreter object
            interpreter = PDFPageInterpreter(manager, device)

            # Start loop processing, processing one page at a time
            for page in pdf_file.get_pages():
                interpreter.process_page(page)
                layout = device.get_result()
                for x in layout:
                    if isinstance(x, LTTextBoxHorizontal):
                        source += x.get_text()
        return source.replace("\n", "。")


def read(path: str) -> str:
    descriptor = path.split(".")[-1]
    if descriptor == "docx":
        w = WordReader(path)
        return w.read()
    elif descriptor == "pdf":
        p = PDFReader(path)
        return p.read()
    elif descriptor == "txt":
        t = TxtRadeer(path)
        return t.read()


if __name__ == '__main__':
    source = read(r"C:\Users\pujic\Desktop\新建文本文档.txt")
    print(source)
